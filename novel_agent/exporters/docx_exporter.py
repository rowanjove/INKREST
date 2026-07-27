"""DOCX publication exporter using python-docx."""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, Optional

from novel_agent.exporters.chapter_export import chapter_heading, collect_publication_book


def export_docx(
    root_dir: Path,
    output_path: Path,
    chapter_ids: Optional[Iterable[str]] = None,
    title: str = "未命名小说",
) -> Path:
    try:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX export; install the production requirements"
        ) from exc

    book = collect_publication_book(
        root_dir,
        title=title,
        chapter_ids=chapter_ids,
    )
    if not book.chapters:
        raise ValueError("No chapters found to export")

    document = docx.Document()
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(11)
    title_paragraph = document.add_heading(book.title, 0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for chapter in book.chapters:
        heading = document.add_heading(chapter_heading(chapter), level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for line in chapter.plain_text.splitlines():
            if not line.strip():
                continue
            paragraph = document.add_paragraph(line.strip())
            paragraph.paragraph_format.first_line_indent = Pt(22)
            paragraph.paragraph_format.line_spacing = 1.5

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    return output
